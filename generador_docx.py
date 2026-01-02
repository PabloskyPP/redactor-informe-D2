"""
Módulo para generar el informe en formato DOCX
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image
from textos import (
    PARRAFOS_FIJOS, TEXTO_DISCREPANCIA_NO_SIGNIFICATIVA, TEXTO_DISCREPANCIA_SIGNIFICATIVA, PARRAFOS_PD
)

def agregar_portada(doc: Document, nombre_completo: str, datos: dict) -> None:
    """
    Agrega la portada del informe
    
    Args:
        doc: Documento de Word
        nombre_completo: Nombre completo del encuestado
        datos: Diccionario con datos generales (edad, fecha_aplicacion, etc.)
    """
    # Título
    titulo = doc.add_heading('TEST DE MATRICES PROGRESIVAS DE RAVEN', 0)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Aumentar tamaño de fuente del título
    for run in titulo.runs:
        run.font.size = Pt(24)
    
    # Espacio
    doc.add_paragraph().add_run().add_break()
    
    # Información del participante
    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Nombre del encuestado
    nombre_run = info.add_run(f"Nombre del encuestado: {nombre_completo}\n")
    nombre_run.bold = True
    nombre_run.font.size = Pt(14)
    
    # Edad
    if datos.get('edad'):
        edad_run = info.add_run(f"Edad: {datos['edad']} años\n")
        edad_run.font.size = Pt(14)
    
    # Fecha de aplicación (si está disponible)
    if datos.get('fecha_aplicacion'):
        fecha_app_run = info.add_run(f"Fecha de aplicación: {datos['fecha_aplicacion']}\n")
        fecha_app_run.font.size = Pt(14)
    
    # Fecha del informe
    fecha_actual = datetime.now().strftime("%d/%m/%Y")
    fecha_informe_run = info.add_run(f"Fecha del informe: {fecha_actual}\n")
    fecha_informe_run.font.size = Pt(14)
    
    # Espacio
    doc.add_paragraph().add_run().add_break()
    
    # Nota confidencial
    nota = doc.add_paragraph()
    nota.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    nota_run = nota.add_run(f"Este es un informe de evaluación cognitiva, obtenido a partir del rendimiento de {nombre_completo} en la prueba Matrices Progresivas de Raven Escala Estándar.")
    nota.add_run("\nInforme confidencial de uso profesional y educativo")
    nota_run.italic = True
    nota_run.font.size = Pt(12)

def crear_informe_docx(resultados, clasificaciones, nombre_caso="caso"):
    """
    Crea el documento Word con el informe del test D2
    
    Args:
        resultados: Dict con puntuaciones directas y datos del encuestado
        clasificaciones: Dict con clasificaciones (PT)
        nombre_caso: Nombre del evaluado (fallback si no está en resultados)
        
    Returns:
        Document: Documento Word generado
    """
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Extraer nombres de resultados
    nombre_completo = resultados.get('nombre_completo') or nombre_caso
    nombre = resultados.get('nombre') or nombre_caso
    
    # Preparar datos para la portada
    datos_portada = {
        'edad': resultados.get('edad'),
        'fecha_aplicacion': resultados.get('fecha_aplicacion')
    }
    
    # 1. PORTADA (Primera página)
    agregar_portada(doc, nombre_completo, datos_portada)
    doc.add_page_break()
    
    # ========================================================================
    # TÍTULO E INTRODUCCIÓN
    # ========================================================================
    # ========================================================================
    # DESCRIPCIÓN DE LA PRUEBA
    # ========================================================================
    subtitulo = doc.add_paragraph()
    run = subtitulo.add_run("Descripción de la prueba")
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph(PARRAFOS_FIJOS['descripcion_prueba'])
    
    # ========================================================================
    # VARIABLES DEL TEST (Descripción técnica)
    # ========================================================================
    doc.add_paragraph("A continuación, se presenta una relación de las puntuaciones o variables de la prueba; se especifica cómo se obtienen y se dan algunas características de su significación psicológica y psicométrica:")


    
    # ========================================================================
    # INSERTAR IMAGEN G1 baremos
    # ========================================================================
    
    # Título de resultados
    titulo_resultados = doc.add_paragraph()
    titulo_resultados.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = titulo_resultados.add_run(PARRAFOS_FIJOS['titulo_resultados'])
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()  # Espacio

    # Introducción a la imagen de baremos
    subtitulo = doc.add_paragraph()
    run = subtitulo.add_run("texto_imagen_baremos")
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph(PARRAFOS_FIJOS['texto_imagen_baremos'])

    # Ruta al gráfico final (en el mismo directorio que el script)
    # Ahora usamos grafico_D2_final.png que incluye todas las superposiciones
    script_dir = os.path.dirname(os.path.abspath(__file__))
    grafico_path = os.path.join(script_dir, 'baremos_Raven.png')
    
    if os.path.exists(grafico_path):
        # Añadir el gráfico con dimensionamiento mejorado
        # Obtener dimensiones de página (ignorar márgenes como se especifica en issue)
        section = doc.sections[-1]
        page_width = section.page_width
        page_height = section.page_height
        
        # Leer dimensiones de la imagen
        try:
            img = Image.open(grafico_path)
            img_width, img_height = img.size
            img_aspect = img_width / img_height
            page_aspect = page_width / page_height
            
            # Calcular dimensiones finales para ajustar tanto al techo como a la base
            # manteniendo la proporción y sin deformación
            if img_aspect > page_aspect:
                # La imagen es más ancha proporcionalmente - limitar por ancho
                final_width = page_width
                final_height = page_width / img_aspect
            else:
                # La imagen es más alta proporcionalmente - limitar por alto
                final_height = page_height
                final_width = page_height * img_aspect
            
            # Añadir la imagen al documento con el tamaño calculado
            paragraph_img = doc.add_paragraph()
            paragraph_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run_img = paragraph_img.add_run()
            picture = run_img.add_picture(grafico_path, width=int(final_width), height=int(final_height))
            
        except Exception as e:
            # Si hay algún error, usar un tamaño fijo razonable
            paragraph_img = doc.add_paragraph()
            paragraph_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run_img = paragraph_img.add_run()
            run_img.add_picture(grafico_path, width=Inches(7))
    
    # Tabla calculo índices de discrepancia
    subtitulo = doc.add_paragraph()
    run = subtitulo.add_run("texto_tabla_discrepancia")
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph(PARRAFOS_FIJOS['texto_tabla_discrepancia'])

    # Tabla con puntuaciones directas obtenidas y esperadas y resto de ambas como índice de discrepancia
    tabla = doc.add_table(rows=3, cols=8)
    tabla.style = 'Light Grid Accent 1'

    # Encabezado de párrafo para el índice de discrepancia
    encabezado_discrepancia = doc.add_paragraph()
    run = encabezado_discrepancia.add_run("Resultado respecto al índice de discrepancia")
    run.bold = True
    run.font.size = Pt(12)

    # ========================================================================
    # PÁRRAFOS CONDICIONALES
    # ========================================================================
    
    # Párrafo conidional según discrepancia significativa o no
    if resultados.get('Discrepancia_significativa', False):
        doc.add_paragraph(TEXTO_DISCREPANCIA_SIGNIFICATIVA.format(nombre=nombre))
    else:
        doc.add_paragraph(TEXTO_DISCREPANCIA_NO_SIGNIFICATIVA.format(nombre=nombre))
    doc.add_paragraph()  # Espacio

    # ========================================================================
    # DESCRIPCIÓN DE PUNTUACIONES DIRECTAS

    # Encabezado de párrafo para el índice de discrepancia
    encabezado_discrepancia = doc.add_paragraph()
    run = encabezado_discrepancia.add_run("Resultado respecto a la puntuación directa de la prueba")
    run.bold = True
    run.font.size = Pt(12)

    p_pd_titulo = doc.add_paragraph()
    run = p_pd_titulo.add_run("🔹 Puntuaciones directas (PD)"
                            )
    run.bold = True
    run.font.size = Pt(12) 
    
    
    return doc


def guardar_informe(doc, ruta_salida):
    """
    Guarda el documento generado
    
    Args:
        doc: Documento Word
        ruta_salida: Ruta donde guardar el archivo
    """
    doc.save(ruta_salida)
    print(f"Informe generado exitosamente en: {ruta_salida}")
